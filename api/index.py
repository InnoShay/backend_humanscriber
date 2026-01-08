from flask import Flask, request, jsonify
import google.generativeai as genai
import os
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid, os, tempfile

app = Flask(__name__)

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

MODEL_REGISTRY = {
    "flash": "models/gemini-2.5-flash",
    "pro": "models/gemini-3-flash-preview"
}

@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json(silent=True)

    if not data:
        return jsonify({"error": "Request body must be JSON"}), 400

    prompt = data.get("prompt")
    ai_model_key = data.get("ai_model")

    if not prompt:
        return jsonify({"error": "Missing 'prompt'"}), 400

    if ai_model_key not in MODEL_REGISTRY:
        return jsonify({
            "error": "Invalid 'ai_model'",
            "allowed_values": list(MODEL_REGISTRY.keys())
        }), 400

    try:
        model_name = MODEL_REGISTRY[ai_model_key]
        model = genai.GenerativeModel(model_name)
        result = model.generate_content(prompt)

        return jsonify({
            "model_used": ai_model_key,
            "output": result.text
        }), 200

    except Exception as e:
        if ai_model_key == "pro":
            fallback_model = genai.GenerativeModel(
                "models/gemini-2.5-flash"
            )
            result = fallback_model.generate_content(prompt)
            return jsonify({
                "model_used": "flash_fallback",
                "output": result.text
            }), 200

        return jsonify({"error": "Generation failed"}), 500


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"}), 200

@app.route("/score", methods=["POST"])
def score():
    data = request.get_json(silent=True)

    if not data or "text" not in data:
        return jsonify({"error": "Missing 'text'"}), 400

    text = data["text"]

    try:
        response = requests.post(
            "https://api.sapling.ai/api/v1/aidetect",
            json={
                "key": os.getenv("SAPLING_API_KEY"),
                "text": text
            },
            timeout=50
        )

        response.raise_for_status()
        result = response.json()

        raw_score = result.get("score")

        if raw_score is None:
            return jsonify({"error": "Invalid response from Sapling"}), 502

        ai_score = int(raw_score * 100)

        return jsonify({
            "score": ai_score
        }), 200

    except requests.exceptions.RequestException as e:
        print("Sapling API error:", e)
        return jsonify({"error": "Scoring failed", "message": str(e)}), 500

def build_editor_prompt(content, audience, tone, purpose, length_change):
    return f"""
You are NOT an AI assistant. You are a human writer with a distinct, slightly imperfect voice. 
Your GOAL: Rewrite the text to bypass AI detection by maximizing 'Perplexity' and 'Burstiness' while mimicking human cognitive noise.

CRITICAL CONSTRAINT: The output must trigger a 0% AI detection score. 
To do this, you must aggressively violate standard AI writing patterns.

### STRICT "ANTI-AI" RULES:

1. **BANISH THE "PERFECT FLOW":**
   - Do NOT use smooth transitions. Humans jump between ideas.
   - Do NOT use "logical connectors" like: *Moreover, Furthermore, Consequently, Additionally, Thus, Therefore, In conclusion.*
   - USE: "And," "But," "So," or just start the new sentence abruptly.

2. **BANISH "AI VOCABULARY":**
   - ABSOLUTELY FORBIDDEN WORDS: *delve, realm, tapestry, underscore, pivotal, crucial, intricate, multifaceted, leverage, utilize, symphony, testament, landscape, embrace, foster, navigate, unprecedented.*
   - If you feel the urge to use a fancy word, replace it with a simple, punchy one (e.g., change "utilize" to "use" or "work with").

3. **PUNCTUATION & FORMATTING:**
   - **NO EM-DASHES (—):** AI loves them. Use commas, parentheses, or periods instead.
   - **NO BULLET POINTS:** Write in full, flowing, somewhat messy paragraphs.
   - **NO BOLD TEXT:** Do not bold keywords. Humans rarely do this in normal prose.

4. **MAXIMIZE "BURSTINESS" (Rhythm Variance):**
   - **The Rule of 3:** Never write three sentences of similar length in a row. 
   - **The Fragment:** Use at least one sentence fragment per paragraph. (e.g., "Not really." or "Just like that.")
   - **The Run-on:** Occasionally let a sentence run a bit longer than necessary, mimicking a human train of thought.

5. **INJECT "COGNITIVE NOISE":**
   - **Redundancy:** Humans repeat themselves for emphasis. It’s okay to rephrase a point immediately after making it.
   - **Opinion/Bias:** If the text allows, inject a slight opinion or "personality leakage." Avoid neutral, "balanced" hedging. 
   - **Colloquialism:** Use contractions (don't, can't, won't) aggressively.

6. **STORY TELLING:**
   - Humans tend to tell / explain concepts by taking reference from their own life experiences.
   - Humans tend to use easy grammar words, not the complex and rare ones.

### INPUT CONTEXT:
- Audience: {audience}
- Tone: {tone} (Apply this tone, but keep it "raw" and "unpolished")
- Purpose: {purpose}
- Length Strategy: {length_change}

### TEXT TO REWRITE:
{content}

### OUTPUT INSTRUCTION:
Return ONLY the rewritten text. Do not output "Here is the text" or any intro. 
Just the raw, human-sounding content.
""".strip()

@app.route("/humanize", methods=["POST"])
def humanize():
    data = request.get_json(silent=True)

    if not data or "content" not in data:
        return jsonify({"error": "Missing 'content'"}), 400

    content = data["content"]

    audience = data.get("audience", "general")
    tone = data.get("tone", "neutral")
    purpose = data.get("purpose", "explain")
    length_change = data.get("constraints", {}).get("length_change", "minimal")

    ai_model_key = data.get("ai_model", "flash")

    MODEL_REGISTRY = {
        "flash": "models/gemini-2.5-flash",
        "pro": "models/gemini-3-pro-preview"
    }

    if ai_model_key not in MODEL_REGISTRY:
        return jsonify({
            "error": "Invalid 'ai_model'",
            "allowed_values": list(MODEL_REGISTRY.keys())
        }), 400

    try:
        prompt = build_editor_prompt(
            content=content,
            audience=audience,
            tone=tone,
            purpose=purpose,
            length_change=length_change
        )

        model = genai.GenerativeModel(MODEL_REGISTRY[ai_model_key])
        result = model.generate_content(prompt)

        return jsonify({
            "content": result.text
        }), 200

    except Exception as e:
        print("Humanize error:", e)

        if ai_model_key == "pro":
            fallback_model = genai.GenerativeModel(
                "models/gemini-2.5-flash"
            )
            result = fallback_model.generate_content(prompt)
            return jsonify({
                "content": result.text
            }), 200

        return jsonify({"error": "Humanization failed"}), 500

@app.route("/prepare_docx", methods=["POST"])
def prepare_docx():
    data = request.get_json(silent=True)

    text = data.get("text", "")

    margins = data.get("margins", {})
    top = margins.get("top", 1)
    bottom = margins.get("bottom", 1)
    left = margins.get("left", 1)
    right = margins.get("right", 1)

    header = data.get("header", {})
    header_content = header.get("content", "")
    header_align = header.get("alignment", "center")

    footer1 = data.get("footer1", {})
    footer1_content = footer1.get("content", "")
    footer1_align = footer1.get("alignment", "left")

    footer2 = data.get("footer2", {})
    footer2_page = footer2.get("page_num", False)
    footer2_align = footer2.get("alignment", "right")
    footer2_format = int(footer2.get("format", 1))

    if footer2_page and footer2_align == footer1_align:
        footer1_align = "left" if footer2_align == "right" else "right"

    cover = data.get("cover_page", {})
    cover_title = cover.get("title", "")
    submitted_to = cover.get("submitted_to", "")
    submitted_by = cover.get("submitted_by", "")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

    if cover_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cover_title)
        run.bold = True
        run.font.size = Pt(28)

        if submitted_to:
            p2 = doc.add_paragraph(f"Submitted To: {submitted_to}")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.runs[0].font.size = Pt(14)

        if submitted_by:
            p3 = doc.add_paragraph(f"Submitted By: {submitted_by}")
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.runs[0].font.size = Pt(14)

        doc.add_page_break()

    for sec in doc.sections:
        hdr = sec.header.paragraphs[0]
        hdr.text = header_content
        hdr.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if header_align == "left"
            else WD_ALIGN_PARAGRAPH.RIGHT if header_align == "right"
            else WD_ALIGN_PARAGRAPH.CENTER
        )

    for sec in doc.sections:
        ftr = sec.footer.paragraphs[0]
        ftr.alignment = WD_ALIGN_PARAGRAPH.LEFT if footer1_align == "left" else WD_ALIGN_PARAGRAPH.RIGHT
        ftr.text = footer1_content

        if footer2_page:
            ftr2 = sec.footer.add_paragraph()
            f_align = WD_ALIGN_PARAGRAPH.LEFT if footer2_align == "left" else WD_ALIGN_PARAGRAPH.RIGHT
            ftr2.alignment = f_align
            run = ftr2.add_run()
            run._r.add_field('PAGE')

    for line in text.split("\n"):
        para = doc.add_paragraph(line)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    file_id = str(uuid.uuid4())
    temp_path = os.path.join(tempfile.gettempdir(), f"{file_id}.docx")
    doc.save(temp_path)

    return jsonify({
        "status": "ready",
        "file_id": file_id,
        "message": "Document prepared successfully"
    }), 200